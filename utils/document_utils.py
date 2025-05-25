import fitz  # PyMuPDF
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.tree import DecisionTreeClassifier

# ---------- استخراج العنوان ----------

def extract_title_pdf(file_path):
    doc = fitz.open(file_path)
    title = None
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text:
            title = text.split("\n")[0]
            break
    return title

def extract_title_docx(file_path):
    doc = docx.Document(file_path)
    return doc.paragraphs[0].text if doc.paragraphs else ""

# ---------- البحث داخل المستند ----------

def search_pdf(file_path, search_text):
    doc = fitz.open(file_path)
    results = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if search_text.lower() in text.lower():
            results.append(page_num)
    return results

def search_docx(file_path, search_text):
    doc = docx.Document(file_path)
    results = []
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            results.append(i)
    return results

# ---------- استخراج النص الكامل ----------

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text("text")
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

# ---------- التصنيف ----------

def classify_document(text):
    # أمثلة بسيطة لتدريب المصنف
    examples = [
        ("هذا تقرير مالي يحتوي على معلومات مالية وتحليل بيانات", "تقرير مالي"),
        ("تقرير طبي عن حالة مريض", "تقرير طبي"),
        ("مقالة أدبية عن الشعر العربي", "أدب"),
        ("بيانات مبيعات الشركة السنوية", "تقرير مالي"),
        ("تشخيص وعلاج مرض السكري", "تقرير طبي"),
        ("قصيدة نثرية معبرة", "أدب"),
    ]
    contents = [e[0] for e in examples]
    labels = [e[1] for e in examples]

    vectorizer = TfidfVectorizer()
    X = vectorizer.fit_transform(contents)
    classifier = DecisionTreeClassifier()
    classifier.fit(X, labels)

    input_vector = vectorizer.transform([text])
    prediction = classifier.predict(input_vector)
    return prediction[0]

# ---------- المعاينة ----------

def preview_document(file_path):
    content = ""
    if file_path.lower().endswith('.pdf'):
        doc = fitz.open(file_path)
        for page in doc:
            content += page.get_text("text")
            if len(content) > 1000:
                break
    elif file_path.lower().endswith('.docx'):
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            content += para.text + "\n"
            if len(content) > 1000:
                break
    return content[:1000]
